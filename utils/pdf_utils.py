# PDF utility functions for merging, splitting, and converting PDFs
import os
import zipfile
from pypdf import PdfReader, PdfWriter
from pdf2image import convert_from_path
from PIL import Image
import tempfile

def merge_pdfs(pdf_files, output_path="merged.pdf"):
    """
    Merge multiple PDF files into one.
    
    Args:
        pdf_files: List of file paths to PDF files
        output_path: Output file path for merged PDF
    
    Returns:
        str: Path to the merged PDF file
    """
    writer = PdfWriter()
    
    for pdf_file in pdf_files:
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            writer.add_page(page)
    
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    
    return output_path

def split_pdf(pdf_file, page_range, output_path="split.pdf"):
    """
    Split a PDF file by extracting specific pages.
    
    Args:
        pdf_file: Path to the input PDF file
        page_range: String like "5-8" or "3" for page range or single page
        output_path: Output file path for split PDF
    
    Returns:
        str: Path to the split PDF file
    """
    reader = PdfReader(pdf_file)
    writer = PdfWriter()
    
    # Parse page range
    if '-' in page_range:
        start_page, end_page = map(int, page_range.split('-'))
        start_page -= 1  # Convert to 0-based indexing
        end_page -= 1
    else:
        start_page = int(page_range) - 1  # Convert to 0-based indexing
        end_page = start_page
    
    # Add pages to writer
    for page_num in range(start_page, end_page + 1):
        if page_num < len(reader.pages):
            writer.add_page(reader.pages[page_num])
    
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    
    return output_path

def pdf_to_images(pdf_file, output_dir="images", format="PNG"):
    """
    Convert PDF pages to images.
    
    Args:
        pdf_file: Path to the input PDF file
        output_dir: Directory to save images
        format: Image format (PNG or JPEG)
    
    Returns:
        list: List of image file paths
    """
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Convert PDF to images
        images = convert_from_path(pdf_file)
        image_paths = []
        
        for i, image in enumerate(images):
            image_path = os.path.join(output_dir, f"page_{i+1}.{format.lower()}")
            image.save(image_path, format)
            image_paths.append(image_path)
        
        return image_paths
    
    except Exception as e:
        # Handle poppler not installed error
        if "poppler" in str(e).lower():
            raise Exception("Poppler is required for PDF to image conversion. "
                          "Please install poppler-utils (Linux) or poppler (Windows/Mac)")
        else:
            raise e

def create_zip_from_images(image_paths, zip_path="images.zip"):
    """
    Create a ZIP file containing all images.
    
    Args:
        image_paths: List of image file paths
        zip_path: Output ZIP file path
    
    Returns:
        str: Path to the ZIP file
    """
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for image_path in image_paths:
            zipf.write(image_path, os.path.basename(image_path))
    
    return zip_path

def word_to_pdf(word_file, output_path="converted.pdf"):
    """
    Convert Word document to PDF preserving formatting using Linux-compatible libraries.
    
    Args:
        word_file: Path to the input Word file (.docx, .doc)
        output_path: Output file path for PDF
    
    Returns:
        str: Path to the converted PDF file
    """
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from PIL import Image
        import zipfile
        import io
        
        # Read the Word document
        doc = Document(word_file)
        
        # Create PDF with better margins
        pdf_doc = SimpleDocTemplate(
            output_path, 
            pagesize=A4,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Create custom styles for different formatting
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.black,
            alignment=TA_LEFT
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            textColor=colors.black,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            textColor=colors.black,
            alignment=TA_LEFT,
            leftIndent=0,
            rightIndent=0
        )
        
        # Process paragraphs with formatting
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Detect paragraph style
                para_style = normal_style
                
                # Check for heading styles
                if paragraph.style.name.startswith('Heading'):
                    if '1' in paragraph.style.name:
                        para_style = title_style
                    else:
                        para_style = heading_style
                
                # Process runs to preserve bold/italic
                formatted_text = ""
                for run in paragraph.runs:
                    text = run.text
                    if text:
                        # Apply formatting tags
                        if run.bold and run.italic:
                            text = f"<b><i>{text}</i></b>"
                        elif run.bold:
                            text = f"<b>{text}</b>"
                        elif run.italic:
                            text = f"<i>{text}</i>"
                        elif run.underline:
                            text = f"<u>{text}</u>"
                        
                        formatted_text += text
                
                if formatted_text.strip():
                    # Handle different alignments
                    if paragraph.alignment == 1:  # Center
                        para_style.alignment = TA_CENTER
                    elif paragraph.alignment == 2:  # Right
                        para_style.alignment = TA_RIGHT
                    elif paragraph.alignment == 3:  # Justify
                        para_style.alignment = TA_JUSTIFY
                    else:
                        para_style.alignment = TA_LEFT
                    
                    p = Paragraph(formatted_text, para_style)
                    story.append(p)
                    story.append(Spacer(1, 6))
        
        # Process images from the document
        image_count = 0
        try:
            # Extract all images from the Word document
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_count += 1
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Create temporary image file
                        temp_image = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                        temp_image.write(image_data)
                        temp_image.close()
                        
                        # Add image to PDF
                        try:
                            # Get image dimensions and resize if needed
                            with Image.open(temp_image.name) as img:
                                img_width, img_height = img.size
                                
                                # Calculate size to fit in page (max 6 inches wide)
                                max_width = 6 * inch
                                max_height = 4 * inch
                                
                                # Maintain aspect ratio
                                if img_width > max_width:
                                    ratio = max_width / img_width
                                    new_width = max_width
                                    new_height = img_height * ratio
                                else:
                                    new_width = min(img_width * 0.75, max_width)  # Convert pixels to points
                                    new_height = min(img_height * 0.75, max_height)
                                
                                if new_height > max_height:
                                    ratio = max_height / new_height
                                    new_width *= ratio
                                    new_height = max_height
                            
                            # Add image title
                            if image_count == 1:
                                story.append(Spacer(1, 12))
                                img_title = Paragraph("<b>Images from Document:</b>", normal_style)
                                story.append(img_title)
                                story.append(Spacer(1, 6))
                            
                            # Add image to story
                            rl_image = RLImage(temp_image.name, width=new_width, height=new_height)
                            story.append(rl_image)
                            story.append(Spacer(1, 12))
                            
                        except Exception as img_error:
                            print(f"Warning: Could not process image {image_count}: {img_error}")
                        
                        # Clean up temp file
                        try:
                            os.unlink(temp_image.name)
                        except:
                            pass
                            
                    except Exception as e:
                        print(f"Warning: Could not extract image {image_count}: {e}")
        except Exception as e:
            print(f"Warning: Could not process images: {e}")
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text with basic formatting
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            if text:
                                if run.bold:
                                    text = f"<b>{text}</b>"
                                elif run.italic:
                                    text = f"<i>{text}</i>"
                                cell_text += text
                    row_data.append(cell_text or " ")
                table_data.append(row_data)
            
            if table_data:
                # Create table with styling
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
        
        # Build the PDF
        pdf_doc.build(story)
        return output_path
        
    except ImportError as e:
        raise Exception(f"Word to PDF conversion requires python-docx and reportlab packages: {e}")
    except Exception as e:
        raise Exception(f"Error converting Word document: {e}")
